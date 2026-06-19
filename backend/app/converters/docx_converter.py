"""DOCX to PDF, TXT, and MD converters."""

from __future__ import annotations

import logging
import shutil
import tempfile
from pathlib import Path

import docx
from markdownify import markdownify as md_from_html

from app.converters.utils import output_path, run_subprocess

logger = logging.getLogger("docforge.converters.docx")


def _find_soffice() -> str:
    """Locate the LibreOffice / soffice binary."""
    candidates = [
        "soffice",
        "libreoffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/bin/soffice",
        "/usr/bin/libreoffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ]
    for candidate in candidates:
        if shutil.which(candidate):
            return candidate
        path_candidate = Path(candidate)
        if path_candidate.exists():
            return str(path_candidate)
    raise RuntimeError(
        "LibreOffice not found. Install it from https://www.libreoffice.org/download/ "
        "and ensure 'soffice' is on your PATH."
    )


def docx_to_pdf(input_file: Path, output_dir: Path) -> Path:
    """Convert a DOCX to PDF using LibreOffice headless."""
    soffice = _find_soffice()
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        staged_input = temp_path / input_file.name
        shutil.copy2(input_file, staged_input)
        run_subprocess(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(temp_path), str(staged_input)],
            context="DOCX→PDF via LibreOffice",
        )
        generated_pdf = temp_path / f"{input_file.stem}.pdf"
        if not generated_pdf.exists():
            raise RuntimeError(f"LibreOffice ran but output PDF not found at {generated_pdf}.")
        out = output_path(output_dir, input_file.stem, "pdf")
        shutil.copy2(generated_pdf, out)
    logger.info("docx→pdf done: %s → %s", input_file.name, out.name)
    return out


def docx_to_txt(input_file: Path, output_dir: Path) -> Path:
    """Extract plain text from a DOCX."""
    out = output_path(output_dir, input_file.stem, "txt")
    document = docx.Document(str(input_file))
    lines: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            lines.append("\t".join(cell.text for cell in row.cells))
    out.write_text("\n".join(lines), encoding="utf-8")
    logger.info("docx→txt done: %s", input_file.name)
    return out


def docx_to_md(input_file: Path, output_dir: Path) -> Path:
    """Convert DOCX to Markdown."""
    out = output_path(output_dir, input_file.stem, "md")
    document = docx.Document(str(input_file))
    html_parts: list[str] = []
    for paragraph in document.paragraphs:
        html_parts.append(f"<p>{paragraph.text}</p>")
    for table in document.tables:
        html_parts.append("<table>")
        for row in table.rows:
            html_parts.append("<tr>" + "".join(f"<td>{cell.text}</td>" for cell in row.cells) + "</tr>")
        html_parts.append("</table>")
    markdown_text = md_from_html("\n".join(html_parts), heading_style="ATX")
    out.write_text(markdown_text, encoding="utf-8")
    logger.info("docx→md done: %s", input_file.name)
    return out
