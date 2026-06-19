"""
DOCX merge — append multiple Word documents into one using python-docx.

Strategy: open the first document, then for each subsequent document
add a page break and copy all paragraphs and tables into the first doc.
"""

import logging
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger("docforge.mergers.docx")


def _add_page_break(doc: Document) -> None:
    """Insert a page break paragraph at the end of a document."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _copy_element(element):
    """Deep-copy an XML element for insertion into another document."""
    return deepcopy(element)


def merge_docx(input_files: list[Path], output_dir: Path) -> Path:
    """
    Merge a list of DOCX files into a single DOCX.

    Opens the first document as the base, then appends each subsequent
    document's body elements after a page break.

    Args:
        input_files: ordered list of DOCX file paths (minimum 2)
        output_dir: directory to write the merged output

    Returns:
        Path to the merged DOCX file.

    Raises:
        ValueError: if fewer than 2 files provided
        RuntimeError: if a file cannot be opened as DOCX
    """
    if len(input_files) < 2:
        raise ValueError("Merge requires at least 2 files.")

    try:
        merged = Document(str(input_files[0]))
        logger.info("merge_docx: base document = %s", input_files[0].name)
    except Exception as exc:
        raise RuntimeError(
            f"Could not open '{input_files[0].name}' as a Word document."
        ) from exc

    for docx_path in input_files[1:]:
        try:
            src = Document(str(docx_path))
        except Exception as exc:
            raise RuntimeError(
                f"Could not open '{docx_path.name}' as a Word document."
            ) from exc

        # Add page break before each appended document
        _add_page_break(merged)

        # Copy body elements (paragraphs + tables) from source into merged
        for element in src.element.body:
            # Skip the final sectPr (section properties) to avoid layout conflicts
            if element.tag.endswith("}sectPr"):
                continue
            merged.element.body.append(_copy_element(element))

        logger.info("merge_docx: appended %s", docx_path.name)

    out = output_dir / "merged.docx"
    merged.save(str(out))
    logger.info("merge_docx: wrote %s (%d files)", out.name, len(input_files))
    return out
